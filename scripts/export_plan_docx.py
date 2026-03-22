#!/usr/bin/env python3
"""
Export plan.md to CRM_Plan.docx for use in Google Docs.

Google Docs does not use a downloadable ".gdoc" format; upload the generated
.docx to Google Drive and choose Open with > Google Docs to get an editable Doc.
"""
from __future__ import annotations

import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
DEFAULT_INPUT = ROOT / "plan.md"
DEFAULT_OUTPUT = ROOT / "CRM_Plan.docx"


def strip_frontmatter(text: str) -> str:
    if text.startswith("---"):
        end = text.find("\n---", 3)
        if end != -1:
            return text[end + 4 :].lstrip("\n")
    return text


def strip_mermaid_blocks(text: str) -> str:
    """Remove fenced mermaid blocks; Word/Google import does not render them."""
    out: list[str] = []
    i = 0
    lines = text.splitlines()
    while i < len(lines):
        line = lines[i]
        if line.strip().startswith("```") and "mermaid" in line.lower():
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                i += 1
            if i < len(lines):
                i += 1
            out.append("\n*[Mermaid diagram omitted — see plan.md in the repo.]*\n")
            continue
        out.append(line)
        i += 1
    return "\n".join(out)


def export_with_pandoc(md_path: Path, out_path: Path) -> bool:
    pandoc = shutil.which("pandoc")
    if not pandoc:
        return False
    subprocess.run(
        [pandoc, str(md_path), "-o", str(out_path), "--from=markdown", "--to=docx"],
        check=True,
        cwd=str(ROOT),
    )
    return True


def export_with_docx_fallback(text: str, out_path: Path) -> None:
    try:
        from docx import Document
    except ImportError as e:
        raise SystemExit(
            "Need pandoc (brew install pandoc) or: pip install python-docx\n" + str(e)
        ) from e

    doc = Document()
    in_fence = False
    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith("```"):
            in_fence = not in_fence
            continue
        if in_fence:
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif stripped == "":
            continue
        else:
            # Strip simple markdown emphasis for readability
            plain = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", line)
            plain = re.sub(r"\*\*([^*]+)\*\*", r"\1", plain)
            plain = re.sub(r"`([^`]+)`", r"\1", plain)
            doc.add_paragraph(plain)

    doc.save(out_path)


def main(argv: list[str]) -> int:
    in_path = Path(argv[1]) if len(argv) > 1 else DEFAULT_INPUT
    out_path = Path(argv[2]) if len(argv) > 2 else DEFAULT_OUTPUT

    if not in_path.is_file():
        print(f"Input not found: {in_path}", file=sys.stderr)
        return 1

    body = strip_frontmatter(in_path.read_text(encoding="utf-8"))
    body = strip_mermaid_blocks(body)

    with tempfile.NamedTemporaryFile(
        mode="w",
        suffix=".md",
        encoding="utf-8",
        delete=False,
        dir=str(ROOT),
    ) as tmp:
        tmp.write(body)
        tmp_path = Path(tmp.name)

    try:
        if export_with_pandoc(tmp_path, out_path):
            print(f"Wrote {out_path} (pandoc)")
            return 0
        export_with_docx_fallback(body, out_path)
        print(f"Wrote {out_path} (python-docx fallback — install pandoc for tables/full Markdown)")
        return 0
    finally:
        tmp_path.unlink(missing_ok=True)


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
